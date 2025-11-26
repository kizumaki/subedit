import streamlit as st
import pandas as pd
import re
import io
import os
import pysrt
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_TAB_ALIGNMENT
import random
from datetime import datetime

# --- GLOBAL CONFIGURATION (Unified) ---
TARGET_FONT = 'Times New Roman'
TARGET_SIZE_PT = Pt(12) # Used by all DOCX functions
MAX_SPEAKER_NAME_LENGTH = 35 
MAX_SPEAKER_NAME_WORDS = 4 

# List of common non-speaker phrases (from srt-excel-converter-app.py)
NON_SPEAKER_PHRASES = [
    "the only problem", "note", "warning", "things", "and on the way we came across this", 
    "this is the highest swing in europe", "and i swear", "which meant", "the only thing is", 
    "and remember", "official distance", "first and foremost", "i said", 
    "here we go", "next up", "step 1", "step 2", "step 3", "and step 3", "first up", 
    "so the question is", "i was growing up", "you might be wondering", "update", 
    "nashville to miami", "all i know is", "unlike judy", "the good news is", 
    "aer lingus seat", "the true test is", "just as i suspected", "like i said", 
    "star review and said", "i told them all", "and best of all", "the point is", 
    "americans", "i was thinking", "and they go", "first of all", "second", 
    "are you like", "as a reminder", "round 2", "round 1", "round 3", "round 4", 
    "round 5", "welcome to round 3", "the question is", "quick reminder", 
    "in 2nd place", "coming up", "first stop", "next step", "and that means", 
    "hashtag", "so to be clear", "your second word", "welcome to round 6", 
    "battle finale time", "number 1", "number 2", "but the truth is", 
    "score to beat", "and your winner", "\"crafty\" and \"betcha\". coming up", 
    "next one", "keep in mind", "and it says", "you could say", "welcome to round 2", 
    "and the best part", "onto round 2", "the ride we chose", "good news is", 
    "bad news", "good news", "he thought", "3 teams remain"
]

# Color palette for DataFrame preview (from srt-excel-converter-app.py)
COLOR_PALETTE = [
    'background-color: #ADD8E6; color: #000000', 'background-color: #90EE90; color: #000000', 
    'background-color: #FFB6C1; color: #000000', 'background-color: #FFFFE0; color: #000000', 
    'background-color: #DDA0DD; color: #000000', 'background-color: #AFEEEE; color: #000000', 
    'background-color: #F0E68C; color: #000000', 'background-color: #FFA07A; color: #000000', 
    'background-color: #E0FFFF; color: #000000', 'background-color: #F5F5DC; color: #000000', 
    'background-color: #2F4F4F; color: #FFFFFF', 'background-color: #191970; color: #FFFFFF', 
    'background-color: #006400; color: #FFFFFF', 'background-color: #800000; color: #FFFFFF', 
    'background-color: #4B0082; color: #FFFFFF', 'background-color: #556B2F; color: #FFFFFF', 
    'background-color: #8B4513; color: #FFFFFF', 'background-color: #36454F; color: #FFFFFF',
]

# Shared Regexes for Word Formatter (from word-editor-app-app.py)
SPEAKER_REGEX = re.compile(r"^([A-Z][a-z\s&]+):\s*", re.IGNORECASE)
TIMECODE_REGEX = re.compile(r"^\d{2}:\d{2}:\d{2},\d{3}\s+-->\s+\d{2}:\d{2}:\d{2},\d{3}$")
HTML_CONTENT_REGEX = re.compile(r"((?:</?[ibu]>)+)(.*?)(?:</?[ibu]>)+", re.IGNORECASE | re.DOTALL) 

# Color variables for Word Formatter (from word-editor-app-app.py)
def generate_vibrant_rgb_colors(count=150):
    """Generates a list of highly saturated, distinct RGB colors."""
    colors = set()
    while len(colors) < count:
        h = random.random(); s = 0.8; v = 0.9 
        if s == 0.0: r = g = b = v
        else:
            i = int(h * 6.0); f = h * 6.0 - i; p = v * (1.0 - s); q = v * (1.0 - s * f); t = v * (1.0 - s * (1.0 - f))
            if i % 6 == 0: r, g, b = v, t, p
            elif i % 6 == 1: r, g, b = q, v, p
            elif i % 6 == 2: r, g, b = p, v, t
            elif i % 6 == 3: r, g, b = p, q, v
            elif i % 6 == 4: r, g, b = t, p, v
            else: r, g, b = v, p, q
        r, g, b = int(r * 255), int(g * 255), int(b * 255)
        if (r < 50 and g < 50 and b < 50) or (r > 200 and g > 200 and b > 200): continue 
        colors.add((r, g, b))
    return list(colors)

FONT_COLORS_RGB_150 = generate_vibrant_rgb_colors(150)
speaker_color_map = {}
used_colors = []

def get_speaker_color(speaker_name):
    """Assigns and retrieves a unique color for a given speaker name (for Word Formatter)."""
    global used_colors
    global speaker_color_map
    
    if speaker_name not in speaker_color_map:
        if not used_colors:
            used_colors = [RGBColor(r, g, b) for r, g, b in FONT_COLORS_RGB_150]
            random.shuffle(used_colors)
            
        color_object = used_colors.pop()
        speaker_color_map[speaker_name] = color_object
        
    return speaker_color_map[speaker_name]

# --- CORE LOGIC FUNCTIONS ---

# 1. SRT TO WORD (from srttowordapp.py)
def set_font_and_size(run, font_name, font_size):
    """Applies Font and Size to a specific run."""
    run.font.name = font_name
    run.font.size = font_size

def process_srt_to_docx(uploaded_file, file_name_without_ext):
    """Reads SRT file and converts it to DOCX with basic formatting."""
    
    srt_content = uploaded_file.getvalue().decode('utf-8')
    subs = pysrt.from_string(srt_content)
    document = Document()
    
    document.add_heading(f"SRT Conversion: {file_name_without_ext}", level=1)

    for sub in subs:
        # Add Index
        p_index = document.add_paragraph(f"{sub.index}")
        set_font_and_size(p_index.runs[0], TARGET_FONT, TARGET_SIZE_PT)
        p_index.paragraph_format.space_after = Pt(0) 

        # Add Timecode
        timecode_str = f"{sub.start} --> {sub.end}"
        p_timecode = document.add_paragraph(timecode_str)
        set_font_and_size(p_timecode.runs[0], TARGET_FONT, TARGET_SIZE_PT)
        p_timecode.paragraph_format.space_after = Pt(0)
        
        # Add Content (cleans up tags using pysrt)
        p_content = document.add_paragraph(sub.text_without_tags)
        if p_content.runs:
            set_font_and_size(p_content.runs[0], TARGET_FONT, TARGET_SIZE_PT)
        p_content.paragraph_format.space_after = Pt(12) 

    modified_file = io.BytesIO()
    document.save(modified_file)
    modified_file.seek(0)
    
    return modified_file

# 2. SRT TO EXCEL (from srt-excel-converter-app.py)

def clean_dialogue_text(text):
    """
    Converts HTML/XML style formatting tags (i, b, u) to text enclosed in parentheses ().
    Removes any other HTML/XML tags.
    """
    text = re.sub(r'<i[^>]*>(.*?)</i[^>]*>', r'(\1)', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<b[^>]*>(.*?)</b[^>]*>', r'(\1)', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<u[^>]*>(.*?)</u[^>]*>', r'(\1)', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<[^>]*>', '', text, flags=re.DOTALL)
    return re.sub(r'\s+', ' ', text).strip()

def is_valid_speaker_tag(tag):
    """Checks if a tag is likely a speaker name using linguistic heuristics."""
    tag = tag.strip()
    if not tag: return False
    if tag.lower() in NON_SPEAKER_PHRASES: return False
    if len(tag) > MAX_SPEAKER_NAME_LENGTH: return False
    normalized_tag = tag.replace(' and ', ' ').replace(' and', '').replace('&', ' ').strip()
    if not normalized_tag: return False
    word_count = len(normalized_tag.split())
    if word_count > MAX_SPEAKER_NAME_WORDS: return False 
    first_word = normalized_tag.split()[0] if normalized_tag.split() else normalized_tag
    if first_word[0].isalpha() and first_word[0].islower(): return False
    return True

def parse_srt(srt_content):
    """
    Parses SRT content to extract Start, End timecodes, Speaker, and Dialogue 
    (from srt-excel-converter-app.py - complex parser)
    """
    data = []
    blocks = re.split(r'\n\s*\n', srt_content.strip())
    last_known_speaker = "Unknown" 

    def append_row_and_update_state(speaker, dialogue):
        nonlocal last_known_speaker
        # IMPORTANT: Use clean_dialogue_text to remove tags for Excel output
        data.append([time_start, time_end, speaker, clean_dialogue_text(dialogue)]) 
        last_known_speaker = speaker 

    for block in blocks:
        lines = block.strip().split('\n')
        if len(lines) < 3: continue

        time_line = lines[1].strip()
        time_match = re.match(r'(\d{2}:\d{2}:\d{2},\d{3}) --> (\d{2}:\d{2}:\d{2},\d{3})', time_line)
        if not time_match: continue

        time_start = time_match.group(1) 
        time_end = time_match.group(2)   

        dialogue_lines = lines[2:]
        current_dialogue = ""
        block_initial_speaker = last_known_speaker
        
        for line in dialogue_lines:
            line = line.strip()
            if not line: continue

            segments = re.split(r'((?:[\w\s&]+?): )', line)
            
            i = 0
            while i < len(segments):
                segment = segments[i].strip()
                i += 1
                
                if not segment: continue

                if segment.endswith(':') and len(segment) > 1:
                    speaker_tag = segment[:-1].strip()
                    
                    if is_valid_speaker_tag(speaker_tag):
                        
                        if current_dialogue:
                            speaker_to_use = block_initial_speaker if not data or data[-1][0] != time_start else last_known_speaker
                            append_row_and_update_state(speaker_to_use, current_dialogue)
                            current_dialogue = "" 
                            
                        speaker = speaker_tag
                        dialogue_segment = segments[i].strip() if i < len(segments) else ""
                        i += 1
                        
                        if dialogue_segment:
                            append_row_and_update_state(speaker, dialogue_segment)
                            
                        if block_initial_speaker == last_known_speaker:
                             block_initial_speaker = speaker
                            
                    else:
                        dialogue_segment = segments[i].strip() if i < len(segments) else ""
                        i += 1
                        recombined_text = segment + " " + dialogue_segment
                        
                        if current_dialogue: current_dialogue += " " + recombined_text
                        else: current_dialogue = recombined_text
                        
                else:
                    if current_dialogue: current_dialogue += " " + segment
                    else: current_dialogue = segment

        if current_dialogue:
            speaker_to_use = block_initial_speaker if not data or data[-1][0] != time_start else last_known_speaker
            append_row_and_update_state(speaker_to_use, current_dialogue)

    return pd.DataFrame(data, columns=['Start', 'End', 'Speaker', 'Dialogue'])

def apply_styles(df):
    """Applies distinct background color styling and text color per speaker for DataFrame preview (for Excel page)."""
    unique_speakers = df['Speaker'].unique()
    color_map = {
        speaker: COLOR_PALETTE[i % len(COLOR_PALETTE)]
        for i, speaker in enumerate(unique_speakers)
    }

    def highlight_speaker(row):
        color_style = color_map.get(row['Speaker'], 'background-color: #FFFFFF; color: #000000')
        return [color_style] * len(row)
    
    try:
        # Returns a Styler object for display (Original behavior)
        return df.style.apply(highlight_speaker, axis=1)
    except Exception:
        # Fallback in case of Styler issues
        return df

# 3. WORD SCRIPT FORMATTER (from word-editor-app-app.py)

def set_all_text_formatting(doc):
    """Applies Times New Roman 12pt and specific Spacing (Before: 0pt, After: 6pt, Single Line) to all runs/paragraphs."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = TARGET_FONT
            run.font.size = TARGET_SIZE_PT
        
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(6)

def process_docx(uploaded_file, file_name_without_ext):
    """Performs original advanced document modifications and formatting on a DOCX input."""
    
    global speaker_color_map
    global used_colors
    speaker_color_map = {}
    used_colors = [RGBColor(r, g, b) for r, g, b in FONT_COLORS_RGB_150]
    random.shuffle(used_colors)
    
    # Load the DOCX file uploaded by the user
    original_document = Document(io.BytesIO(uploaded_file.getvalue()))
    raw_paragraphs = [p for p in original_document.paragraphs if p.text.strip()]
    
    document = Document()
    
    # --- A. Set Main Title (25pt, 2 blank lines after) ---
    title_paragraph = document.add_paragraph(file_name_without_ext.upper())
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_before = Pt(0)
    title_paragraph.paragraph_format.space_after = Pt(0) 
    
    title_run = title_paragraph.runs[0]
    title_run.font.name = TARGET_FONT
    title_run.font.size = Pt(25) 
    title_run.bold = True
    
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    document.add_paragraph().paragraph_format.space_after = Pt(0)

    # --- B. Process raw paragraphs and add to new document ---
    
    for paragraph in raw_paragraphs:
        text = paragraph.text.strip()
        
        # B.1 Remove SRT Line Numbers (from basic converters)
        if re.fullmatch(r"^\s*\d+\s*$", text):
            continue 
            
        new_paragraph = document.add_paragraph()
        new_paragraph.style = document.styles['Normal']
        new_paragraph.paragraph_format.space_before = Pt(0) 
        new_paragraph.paragraph_format.space_after = Pt(6) 
        
        # B.2 Bold Timecode (Override Space After = 0)
        if TIMECODE_REGEX.match(text):
            new_paragraph.text = text
            for run in new_paragraph.runs:
                run.font.bold = True
            new_paragraph.paragraph_format.space_after = Pt(0) 

        # B.3 Content (Speaker/Content)
        else:
            
            speaker_match = SPEAKER_REGEX.match(text)
            
            if speaker_match:
                # 1. Set Hanging Indent and Tab Stop
                new_paragraph.paragraph_format.left_indent = Inches(1.0)
                new_paragraph.paragraph_format.first_line_indent = Inches(-1.0)
                new_paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(1.0), WD_TAB_ALIGNMENT.LEFT)
                
                speaker_full = speaker_match.group(0) 
                speaker_name = speaker_match.group(1).strip()
                
                font_color_object = get_speaker_color(speaker_name) 
                rest_of_text = text[len(speaker_full):]
                
                # Run for the speaker name (Bold and Font Color)
                run_speaker = new_paragraph.add_run(speaker_full)
                run_speaker.font.bold = True
                run_speaker.font.color.rgb = font_color_object 
                
                # Insert Tab character
                new_paragraph.add_run('\t') 
                
                current_text = rest_of_text
                
            else:
                # No speaker -> No indent
                new_paragraph.paragraph_format.left_indent = None
                new_paragraph.paragraph_format.first_line_indent = None
                current_text = text

            # --- B.4 Process HTML tags within the current_text ---
            
            # Using the original (and potentially buggy) regex for fidelity to the original file
            matches = list(HTML_CONTENT_REGEX.finditer(current_text))
            last_end = 0
            
            if not speaker_match:
                 new_paragraph.text = "" 

            for match in matches:
                tag_text = match.group(2) 
                start, end = match.span()

                # Add text BEFORE the tag (if any)
                if start > last_end:
                    new_paragraph.add_run(current_text[last_end:start])
                
                # Add the HTML content (Bold and Italic)
                run_html = new_paragraph.add_run(tag_text)
                run_html.font.bold = True
                run_html.font.italic = True
                
                last_end = end

            # Add remaining text AFTER the last tag
            if last_end < len(current_text):
                new_paragraph.add_run(current_text[last_end:])
            
            # Handle case with no tag and no speaker (plain text)
            elif not speaker_match and not matches:
                new_paragraph.add_run(current_text)

    # C. Apply General Font/Size and Spacing (Global settings)
    set_all_text_formatting(document)
    
    modified_file = io.BytesIO()
    document.save(modified_file)
    modified_file.seek(0)
    
    return modified_file


# --------------------------------------------------------------------------------
# --- STREAMLIT PAGES ---
# --------------------------------------------------------------------------------

def srt_to_docx_page():
    # UI text from original srttowordapp.py (translated to English)
    st.markdown("## 📄 SRT to Word (.docx) Converter - Basic")
    st.markdown("This function converts the SRT file to DOCX, preserving line structure (index, timecode, content) and formats the output to **Times New Roman, 12pt**.")
    st.markdown("---")

    uploaded_file = st.file_uploader(
        "1. Upload your SRT file (.srt)",
        type=['srt'],
        key="srt_docx_uploader",
        help="Only .srt format is accepted."
    )

    if uploaded_file is not None:
        original_filename = uploaded_file.name
        file_name_without_ext = os.path.splitext(original_filename)[0]
        
        st.info(f"File received: **{original_filename}**.")
        
        if st.button("2. RUN WORD CONVERSION", key="run_srt_docx"):
            with st.spinner('Processing and creating Word file...'):
                try:
                    modified_file_io = process_srt_to_docx(uploaded_file, file_name_without_ext)
                    
                    new_filename = f"CONVERTED_{file_name_without_ext}.docx"

                    st.success("✅ Conversion complete! You can download the file.")
                    
                    st.download_button(
                        label="3. Download Converted Word File",
                        data=modified_file_io,
                        file_name=new_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    st.markdown("---")
                    st.balloons()

                except Exception as e:
                    st.error(f"An error occurred during processing: {e}")
                    st.warning("Please check the format of the input SRT file.")


def srt_to_excel_page():
    # UI text from original srt-excel-converter-app.py (translated to English)
    st.markdown("## 📊 Analyze & Convert SRT to Excel (.xlsx)")
    st.markdown("This function analyzes the SRT file to extract detailed dialogue and corresponding speaker, then exports to an Excel file.")
    st.markdown("---")
    
    # NOTE: Added this warning to manage user expectation about the Excel file format
    st.warning("⚠️ **NOTE ON EXCEL FORMATTING:** The colorful highlighting is for the web preview only and CANNOT be included in the downloaded Excel (.xlsx) file due to format limitations. The downloaded file will contain clean, organized data.")
    st.markdown("---")

    uploaded_file = st.file_uploader("1. Upload your SRT file (.srt)", type="srt", key="srt_excel_uploader")

    if uploaded_file is not None:
        try:
            srt_content = uploaded_file.read().decode("utf-8")
        except Exception:
            st.error("File encoding error. Please ensure your SRT file is correctly encoded (UTF-8 recommended).")
            return

        with st.spinner('Analyzing SRT data...'):
            # Use the original parser (which returns clean text for Excel)
            df_converted = parse_srt(srt_content)
        
        if df_converted.empty:
            st.error("Could not parse any subtitles.")
            return

        st.subheader("📊 Speaker Statistics")
        
        unique_speakers = df_converted['Speaker'].unique()
        actual_speakers = [s for s in unique_speakers if s not in ["Unknown", ""]]
        speaker_count = len(actual_speakers)

        st.success(f"**Total Recognized Speakers:** {speaker_count} people.")
        
        if speaker_count > 0:
            speaker_list_str = ", ".join(actual_speakers)
            st.markdown(f"**List of Speakers:** {speaker_list_str}")
        else:
            st.info("No clear speakers found.")
            
        st.subheader("Converted Data Preview (Web Styling Only)")
        
        # Apply styling ONLY for the web preview
        styled_df_display = apply_styles(df_converted)
        st.dataframe(styled_df_display, use_container_width=True)

        st.markdown("---")
        
        output = io.BytesIO()
        
        # FIX FOR NAMERROR/CORRUPTION: Save the original DataFrame (df_converted) instead of the Styler object
        df_converted.to_excel(output, index=False, engine='openpyxl') 
        output.seek(0)

        original_name_base = uploaded_file.name.rsplit('.', 1)[0]
        file_name = f"{original_name_base}_DATA.xlsx"
        
        st.download_button(
            label="💾 Download Analyzed Excel File (.xlsx)",
            data=output.read(), 
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        st.success(f"File ready for download: **{file_name}**!")


def word_formatter_page():
    # UI text from original word-editor-app-app.py (translated to English)
    st.markdown("## 📝 Automatic Word Script Formatter Tool (DOCX Input)")
    st.markdown("This function automatically formats an existing Word file (.docx) according to professional rules:")
    st.markdown("- **Title:** Uppercase, 25pt size, centered.")
    st.markdown("- **Timecode:** Bold, minimal line spacing.")
    st.markdown("- **Speaker:** Bold, unique color (per speaker), professional hanging indent and tab stop.")
    st.markdown("---")

    uploaded_file = st.file_uploader(
        "1. Upload your Word file (.docx)",
        type=['docx'],
        key="word_formatter_uploader",
        help="Only Microsoft Word .docx format is accepted. **NOTE: This tool is designed to work with the output from the 'SRT to Word (Basic)' tool.**"
    )

    if uploaded_file is not None:
        original_filename = uploaded_file.name
        file_name_without_ext = os.path.splitext(original_filename)[0]
        
        st.info(f"File received: **{original_filename}**.")
        
        if st.button("2. RUN AUTOMATIC FORMATTING", key="run_word_formatter"):
            with st.spinner('Processing and formatting file...'):
                try:
                    # Use the original DOCX processing function
                    modified_file_io = process_docx(uploaded_file, file_name_without_ext)
                    
                    new_filename = f"FORMATTED_{original_filename}"

                    st.success("✅ Formatting complete! You can download the file.")
                    
                    st.download_button(
                        label="3. Download Formatted Word File",
                        data=modified_file_io,
                        file_name=new_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    st.markdown("---")
                    st.balloons()

                except Exception as e:
                    st.error(f"An error occurred during processing: {e}")
                    st.warning("Please check the format of the input file.")


# --------------------------------------------------------------------------------
# --- MAIN APPLICATION ENTRY POINT ---
# --------------------------------------------------------------------------------

def main():
    """Defines the Streamlit application structure using sidebar navigation."""
    
    st.set_page_config(
        page_title="Subtitle & Script Toolkit", 
        layout="wide",
        initial_sidebar_state="expanded"
    )

    st.sidebar.title("🛠️ COMPREHENSIVE TOOLKIT")
    st.sidebar.markdown("Select a function to use:")
    
    # Navigation Radio Buttons
    app_mode = st.sidebar.radio(
        "Function",
        (
            "1. SRT to Word (Basic)",
            "2. SRT to Excel (Analysis)",
            "3. Word Script Formatting"
        )
    )

    st.sidebar.markdown("---")
    st.sidebar.markdown(
        """
        **Usage:**
        - Each function operates independently.
        - Upload the file, run the process, and download the result.
        """
    )
    
    # Route to the selected page function
    if app_mode == "1. SRT to Word (Basic)":
        srt_to_docx_page()
    elif app_mode == "2. SRT to Excel (Analysis)":
        srt_to_excel_page()
    elif app_mode == "3. Word Script Formatting":
        word_formatter_page()

if __name__ == "__main__":
    if not FONT_COLORS_RGB_150:
         FONT_COLORS_RGB_150 = generate_vibrant_rgb_colors(150)
         
    main()
